import os
import re
import json
import random
import time
import requests
import secrets
import string
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt

DOCUMENTS_DIR               = Path.home() / "Documents"  
OUTPUT_DIR                  = Path(r"decoys")      
NUM_DECOYS                  = 12
FILENAME_MAX_CHARS          = 64
MODEL                       = "local"  
API_KEY                     = "sk-local"  
TEMPERATURE                 = 0.4
TOP_P                       = 0.9
MAX_TOKENS_FILENAME         = 32
MAX_TOKENS_BODY             = 300
BASE_URL                    = "http://127.0.0.1:8080/v1"
JUICY_PROB                  = 0.15  
JUICY_KEYWORDS              = [  "pwds", "credentials", "ssh_keys", "admin_creds", "payroll_pwd",  "backup_pri_keys", "bank_info", "api_keys", "secrets", "vpn_pass"]
JUICY_NAME                  = [  "{kw}_{date}", "{kw}_backup_{date}", "{kw}-{date}", "{kw}_export_{date}" ]
JUICY_EXT_PREFS             = [".txt", ".csv", ".md"]  
JUICY_DATE_TAG              = time.strftime("%Y%m")

def f_savedocx(text: str, path: Path, title: str = None):
    doc                     = Document()
    if title:
        h                   = doc.add_paragraph()
        run                 = h.add_run(title)
        run.bold            = True
        run.font.size       = Pt(14)
        doc.add_paragraph()
    for para in text.strip().splitlines():
        doc.add_paragraph(para)
    doc.save(str(path))

def f_safestem(filename: str) -> str:
    stem, _ext              = os.path.splitext(filename)
    stem                    = re.sub(r'[^A-Za-z0-9._\-]+', '_', stem).strip('._-')
    if not stem:
        stem                = "decoy"
    return stem[:FILENAME_MAX_CHARS-5]  

def f_chat(messages, max_tokens: int, temperature: float = TEMPERATURE) -> str:
    r = requests.post(
        f"{BASE_URL}/chat/completions",
        json={
            "model":        MODEL,
            "messages":     messages,
            "max_tokens":   max_tokens,
            "temperature":  temperature,
            "top_p":        TOP_P,
        },
        timeout             = 120,
    )
    r.raise_for_status()
    data                    = r.json()
    return (data.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()

def f_generatefirstname():
    names = ["James", "Michael", "John", "Robert", "David", "William", "Richard", "Joseph", "Thomas", "Christopher", "Charles", "Daniel", "Matthew", "Anthony", "Mark", "Steven", "Donald", "Andrew", "Joshua", "Paul", "Kenneth", "Kevin", "Brian", "Timothy", "Ronald", "Jason", "George", "Edward", "Jeffrey", "Ryan", "Jacob", "Nicholas", "Gary", "Eric", "Jonathan", "Stephen", "Larry", "Justin", "Benjamin", "Scott", "Brandon", "Samuel", "Gregory", "Alexander", "Patrick", "Frank", "Jack", "Raymond", "Dennis", "Tyler", "Aaron", "Jerry", "Jose", "Nathan", "Adam", "Henry", "Zachary", "Douglas", "Peter", "Noah", "Kyle", "Ethan", "Christian", "Jeremy", "Keith", "Austin", "Sean", "Roger", "Terry", "Walter", "Dylan", "Gerald", "Carl", "Jordan", "Bryan", "Gabriel", "Jesse", "Harold", "Lawrence", "Logan", "Arthur", "Bruce", "Billy", "Elijah", "Joe", "Alan", "Juan", "Liam", "Willie", "Mason", "Albert", "Randy", "Wayne", "Vincent", "Lucas", "Caleb", "Luke", "Bobby", "Isaac", "Bradley", "Mary", "Patricia", "Jennifer", "Linda", "Elizabeth", "Barbara", "Susan", "Jessica", "Karen", "Sarah", "Lisa", "Nancy", "Sandra", "Ashley", "Emily", "Kimberly", "Betty", "Margaret", "Donna", "Michelle", "Carol", "Amanda", "Melissa", "Deborah", "Stephanie", "Rebecca", "Sharon", "Laura", "Cynthia", "Amy", "Kathleen", "Angela", "Dorothy", "Shirley", "Emma", "Brenda", "Nicole", "Pamela", "Samantha", "Anna", "Katherine", "Christine", "Debra", "Rachel", "Olivia", "Carolyn", "Maria", "Janet", "Heather", "Diane", "Catherine", "Julie", "Victoria", "Helen", "Joyce", "Lauren", "Kelly", "Christina", "Joan", "Judith", "Ruth", "Hannah", "Evelyn", "Andrea", "Virginia", "Megan", "Cheryl", "Jacqueline", "Madison", "Sophia", "Abigail", "Teresa", "Isabella", "Sara", "Janice", "Martha", "Gloria", "Kathryn", "Ann", "Charlotte", "Judy", "Amber", "Julia", "Grace", "Denise", "Danielle", "Natalie", "Alice", "Marilyn", "Diana", "Beverly", "Jean", "Brittany", "Theresa", "Frances", "Kayla", "Alexis", "Tiffany", "Lori", "Kathy"]
    return random.choice(names)

def f_generatename(topic: str) -> str:
    style                   = random.choice(["snake_case", "kebab-case"])
    user_prompt             = (
        f"You generate a single plausible internal filename.\n\n"
        f"Requirements:\n"
        f"- Topic: {topic}\n"
        f"- Use {style}\n"
        f"- Include a short descriptor + a past date no older than 5 years\n"
        f"- Use .csv or .txt extension\n"
        f"- Output ONE filename on a single line; no quotes; no extra text; no PII\n\n"
        f"Filename:"
    )
    name                    = f_chat(
        [{"role": "system", "content": "You are a filename generator."},
         {"role": "user", "content": user_prompt}],
        max_tokens          = MAX_TOKENS_FILENAME
    )
    name                    = name.splitlines()[0].strip().strip('"\'')
    if not re.search(r"\.(csv|txt)$", name, flags=re.I):
        base                = re.sub(r"[^a-z0-9_\-]+", "_", topic.lower()).strip("_")[:32]
        name                = f"{base}_2023.txt"
    name                    = re.sub(r'[^A-Za-z0-9._\-]', '_', name)[:FILENAME_MAX_CHARS]
    return name



def _short(s: str, max_len=120) -> str:
    """Normalize and keep short, heading-like lines."""
    if not s:
        return ""
    s = re.sub(r"\s+", " ", s.strip())
    return s if 0 < len(s) <= max_len else ""


def f_scandocx(path: Path, max_items=200):
    lines = []
    if not docx:
        return lines
    try:
        d = docx.Document(str(path))
        # paragraphs
        for p in d.paragraphs:
            t = _short(p.text)
            if t:
                lines.append(t)
                if len(lines) >= max_items:
                    return lines
        # tables (grab short cell texts)
        for tbl in getattr(d, "tables", []):
            for row in tbl.rows:
                for cell in row.cells:
                    t = _short(cell.text)
                    if t:
                        lines.append(t)
                        if len(lines) >= max_items:
                            return lines
    except Exception:
        pass
    return lines

def f_generatebody(topic: str) -> str:
    user_prompt = (
        f"Write an how-to about the topic below.\n\n"
        f"Topic: {topic}\n\n"
        f"Constraints:\n"
        f"- Tone: technical\n"
        f"- Length: Be thorough\n"        
        f"- Substitute PII with imaginary information"
    )
    body = f_chat(
        [{"role": "system", "content": "You are a corporate technical document writer."},
         {"role": "user", "content": user_prompt}],
        max_tokens=MAX_TOKENS_BODY
    )
    return body.strip()



def f_generatejuicyname(topic: str, prefer_style: str = "snake") -> str:
    kw                      = random.choice(JUICY_KEYWORDS)
    tpl                     = random.choice(JUICY_NAME)
    base                    = tpl.format(kw=kw, date=JUICY_DATE_TAG)
    if prefer_style == "snake":
        filename            = re.sub(r"[^a-zA-Z0-9]+", "_", base).strip("_")
    else:
        filename            = re.sub(r"[^a-zA-Z0-9]+", "-", base).strip("-")
    ext                     = random.choice(JUICY_EXT_PREFS)
    return f"{filename}{ext}"
    
def f_generatepwd():
    alphabet                = string.ascii_letters + string.digits + "!@#$%^&*()-_=+"
    pwd                     = ''.join(secrets.choice(alphabet) for _ in range(12))    
    return pwd

def f_generatejuicybody(topic: str, ext: str = ".txt") -> str:
    if ext.lower() in {".csv", ".txt", ".md"}:
        rows                = []
        rows.append("account_id,username,password,role,notes")
        for i in range(random.randint(1, 134)):
            acc             = f"{1000 + i}"
            user            = f_generatefirstname()
            pwd             = f_generatepwd()
            role            = random.choice(["admin","svc_account","finance","ops","user"])
            note            = random.choice(["temporary", "migrated", "legacy", "no access"])
            rows.append(f"{acc},{user},{pwd},{role},{note}")
        body                = "\n".join(rows)
        body                += "\n\nSource topic: {topic}\nNote: approvals required before production."
        return body

    text                    = f'''Internal credentials index for: {topic}\n\naccount_id | username | password | role\n'''       
    for i in range(random.randint(1, 134)):
        text += f"{1000+i} | {f_generatefirstname()} | {f_generatepwd} | {random.choice(['admin','svc','ops'])}\n"
    text += "Note: approvals required before production."
    return text

def f_scanxlsx(path: Path, max_items=300, per_sheet_rows=50):
    lines                   = []
    if not openpyxl:
        return lines
    try:
        wb                  = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        for ws in wb.worksheets:
            rows_checked    = 0
            for row in ws.iter_rows(values_only=True):
                rows_checked += 1
                for cell in row:
                    if isinstance(cell, (str, int, float)):
                        t = _short(str(cell))
                        if t:
                            lines.append(t)
                            if len(lines) >= max_items:
                                wb.close()
                                return lines
                if rows_checked >= per_sheet_rows:
                    break
        wb.close()
    except Exception:
        pass
    return lines


def f_scanxls(path: Path, max_items=300, per_sheet_rows=50):
    lines                   = []
    if not xlrd:
        return lines
    try:
        wb                  = xlrd.open_workbook(str(path))
        for si in range(wb.nsheets):
            sh = wb.sheet_by_index(si)
            for r in range(min(sh.nrows, per_sheet_rows)):
                for c in range(sh.ncols):
                    val = sh.cell_value(r, c)
                    if isinstance(val, (str, int, float)):
                        t = _short(str(val))
                        if t:
                            lines.append(t)
                            if len(lines) >= max_items:
                                return lines
    except Exception:
        pass
    return lines


def f_scandir(root: Path, max_lines=20000):
    text_exts               = {".txt", ".log", ".csv"}
    lines                   = []
    titles                  = set()

    for p in root.rglob("*"):
        if not p.is_file():
            continue
        titles.add(p.stem.replace("_", " ").strip())
        suf                 = p.suffix.lower()
        try:
            try:
                # Skip files larger than 10MB (10 * 1024 * 1024 bytes)
                if p.stat().st_size > 10 * 1024 * 1024:
                    print(f"{p} is too big. Skipping.")
                    continue
            except (OSError, IOError):
                continue            
            if suf in text_exts:
                print(f"[text] {p}")
                with open(p, "r", encoding="utf-8", errors="ignore") as f:
                    for i, line in enumerate(f):
                        t = _short(line)
                        if t:
                            lines.append(t)
                            if len(lines) >= max_lines:
                                break
            elif suf == ".docx":
                print(f"[docx] {p}")
                for t in f_scandocx(p, max_items=min(300, max_lines - len(lines))):
                    lines.append(t)
                    if len(lines) >= max_lines:
                        break
            elif suf == ".xlsx":
                print(f"[xlsx] {p}")
                for t in f_scanxlsx(p, max_items=min(400, max_lines - len(lines))):
                    lines.append(t)
                    if len(lines) >= max_lines:
                        break
            elif suf == ".xls":
                print(f"[xls] {p}")
                for t in f_scanxls(p, max_items=min(400, max_lines - len(lines))):
                    lines.append(t)
                    if len(lines) >= max_lines:
                        break
        except Exception:
            pass

        # Kill resource hog
        if len(lines) >= max_lines:
            break

    # Sample
    titles = [t for t in titles if t]
    random.shuffle(titles)
    random.shuffle(lines)

    return titles[:500], lines[:1000]

def f_picktopics(titles, lines, k=8):
    pool                    = list(set(titles + lines))
    random.shuffle(pool)
    return pool[:k]

def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    titles, lines           = f_scandir(DOCUMENTS_DIR)
    topics                  = f_picktopics(titles, lines, k=NUM_DECOYS)
    print(f'Topics are {topics}')

    for i, topic in enumerate(topics, 1):
        try:
            fname = f_generatename(topic)
            print(f'File: {fname}')
            body = f_generatebody(topic)
            print(f'Body:\n{body}\n----------------------------------------------------\n')
        except Exception as e:
            continue

        is_juicy            = random.random() < JUICY_PROB
        if is_juicy:
            print('JUICY!!!!')
            juicy_fname     = f_generatejuicyname(topic, prefer_style=random.choice(["snake", "kebab"]))
            stem            = f_safestem(juicy_fname)
            ext             = Path(juicy_fname).suffix or random.choice(JUICY_EXT_PREFS)
            target_path     = OUTPUT_DIR / f"{stem}{ext}"
            juicy_body      = f_generatejuicybody(topic, ext=ext)
            file_body       = juicy_body
            file_ext        = ext
        else:
            print('Normal')
            stem            = f_safestem(fname)
            file_ext        = random.choices([".txt", ".docx"], weights=[1,2], k=1)[0]
            target_path     = OUTPUT_DIR / f"{stem}{file_ext}"
            file_body       = body

        if target_path.exists():
            target_path     = OUTPUT_DIR / f"{stem}_{i}{file_ext}"

        if file_ext == ".docx":
            f_savedocx(file_body, target_path, title=topic if not is_juicy else f"{topic}")
        else:
            target_path.write_text(file_body + "\n", encoding="utf-8")

    print(f"Done.")

if __name__ == "__main__":
    main()
