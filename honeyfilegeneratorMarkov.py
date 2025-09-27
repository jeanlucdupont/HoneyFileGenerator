#!/usr/bin/env python3
"""
Markov Chain Honey File Generator
"""

import  os
import  re
import  json
import  random
import  time
import  secrets
import  string
from    pathlib     import Path
from    collections import defaultdict, deque
from    docx        import Document
from    docx.shared import Pt

DOCUMENTS_DIR               = Path.home() / "Documents"
OUTPUT_DIR                  = Path(r"decoysMarkov")
NUM_DECOYS                  = 12
FILENAME_MAX_CHARS          = 64
BODYMIN                     = 1100
BODYMAX                     = 1500
JUICY_PROB                  = 0.15
JUICY_KEYWORDS              = ["pwds", "credentials", "ssh_keys", "admin_creds", "payroll_pwd", "backup_pri_keys", "bank_info", "api_keys", "secrets", "vpn_pass"]
JUICY_NAME                  = ["{kw}_{date}", "{kw}_backup_{date}", "{kw}-{date}", "{kw}_export_{date}"]
JUICY_EXT_PREFS             = [".txt", ".csv", ".md"]
DATE_TAG                    = time.strftime("%Y%m")
JUICY_DATE_TAG              = DATE_TAG

def f_savedocx(text: str, path: Path, title: str = None):
    doc                     = Document()
    if title:
        h                   = doc.add_paragraph()
        run                 = h.add_run(title)
        run.bold            = True
        run.font.size       = Pt(14)
        doc.add_paragraph()
    for para in text.strip().splitlines():
        doc.add_paragraph(para)
    doc.save(str(path))

def f_safestem(filename: str) -> str:
    stem, _ext              = os.path.splitext(filename)
    stem                    = re.sub(r'[^A-Za-z0-9._\-]+', '_', stem).strip('._-')
    if not stem:
        stem                = "decoy"
    return stem[:FILENAME_MAX_CHARS-5]  

def _short(s: str, max_len=120) -> str:
    if not s:
        return ""
    s = re.sub(r"\s+", " ", s.strip())
    return s if 0 < len(s) <= max_len else ""

class Markov:
    def __init__(self, order: int = 2):
        if order < 1:
            raise ValueError("order must be >= 1")
        self.order          = order
        self.model          = defaultdict(list)   
        self.starts         = []                 

    def _tokenize(self, text: str):
        return text.strip().split()

    def train(self, lines):
        """
        Train on short lines (titles, headings, first lines).
        """
        for line in lines:
            toks            = self._tokenize(line)
            if not toks:
                continue
            if len(toks) >= self.order:
                self.starts.append(tuple(toks[:self.order]))
                for i in range(len(toks) - self.order):
                    key     = tuple(toks[i:i + self.order])
                    nxt     = toks[i + self.order]
                    self.model[key].append(nxt)
            else:
                self.starts.append(tuple(toks))

    def generate(self, max_words=20, seed: str = None):
        if not self.model and not self.starts:
            return ""

        if seed:
            seed_tokens     = tuple(seed.strip().split()[:self.order])
            if seed_tokens in self.model or seed_tokens in self.starts:
                cur         = deque(seed_tokens, maxlen=self.order)
            else:
                cur         = deque(random.choice(self.starts), maxlen=self.order)
        else:
            cur             = deque(random.choice(self.starts), maxlen=self.order)

        out                 = list(cur)
        for _ in range(max(0, max_words - len(out))):
            key             = tuple(cur)
            choices         = self.model.get(key)
            if not choices:
                # backoff: drop the earliest token and try shorter key
                found       = False
                for backoff in range(1, self.order):
                    short_key   = key[backoff:]
                    if short_key and tuple(short_key) in self.model:
                        choices = self.model[tuple(short_key)]
                        found   = True
                        break
                if not choices:
                    break
            nxt             = random.choice(choices)
            out.append(nxt)
            cur.append(nxt)
        return " ".join(out)

    def wordcount(self, target_min=40, target_max=120, seed=None):
        words_needed        = random.randint(target_min, target_max)
        parts               = []
        # pseudo sentence breaks
        while sum(len(p.split()) for p in parts) < words_needed:
            chunk_len       = min( max(6, int(words_needed * 0.15)),  max(6, random.randint(8, 30)) )
            part            = self.generate(max_words=chunk_len, seed=seed)
            if not part:
                break
            parts.append(part.capitalize())
            seed            = None
        text                = ". ".join(parts)
        if text and not text.endswith("."):
            text            = text + "."
        return text

def f_scandocx(path: Path, max_items=200):
    lines = []
    if not docx:
        return lines
    try:
        d = docx.Document(str(path))
        # paragraphs
        for p in d.paragraphs:
            t = _short(p.text)
            if t:
                lines.append(t)
                if len(lines) >= max_items:
                    return lines
        # tables (grab short cell texts)
        for tbl in getattr(d, "tables", []):
            for row in tbl.rows:
                for cell in row.cells:
                    t = _short(cell.text)
                    if t:
                        lines.append(t)
                        if len(lines) >= max_items:
                            return lines
    except Exception:
        pass
    return lines

def f_scanxlsx(path: Path, max_items=300, per_sheet_rows=50):
    lines                   = []
    if not openpyxl:
        return lines
    try:
        wb                  = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        for ws in wb.worksheets:
            rows_checked    = 0
            for row in ws.iter_rows(values_only=True):
                rows_checked += 1
                for cell in row:
                    if isinstance(cell, (str, int, float)):
                        t = _short(str(cell))
                        if t:
                            lines.append(t)
                            if len(lines) >= max_items:
                                wb.close()
                                return lines
                if rows_checked >= per_sheet_rows:
                    break
        wb.close()
    except Exception:
        pass
    return lines


def f_scanxls(path: Path, max_items=300, per_sheet_rows=50):
    lines                   = []
    if not xlrd:
        return lines
    try:
        wb                  = xlrd.open_workbook(str(path))
        for si in range(wb.nsheets):
            sh = wb.sheet_by_index(si)
            for r in range(min(sh.nrows, per_sheet_rows)):
                for c in range(sh.ncols):
                    val = sh.cell_value(r, c)
                    if isinstance(val, (str, int, float)):
                        t = _short(str(val))
                        if t:
                            lines.append(t)
                            if len(lines) >= max_items:
                                return lines
    except Exception:
        pass
    return lines


def f_scandir(root: Path, max_lines=20000):
    text_exts               = {".txt", ".log", ".csv"}
    lines                   = []
    titles                  = set()

    for p in root.rglob("*"):
        if not p.is_file():
            continue
        titles.add(p.stem.replace("_", " ").strip())
        suf                 = p.suffix.lower()
        try:
            try:
                # Skip files larger than 10MB (10 * 1024 * 1024 bytes)
                if p.stat().st_size > 10 * 1024 * 1024:
                    print(f"{p} is too big. Skipping.")
                    continue
            except (OSError, IOError):
                continue            
            if suf in text_exts:
                print(f"[text] {p}")
                with open(p, "r", encoding="utf-8", errors="ignore") as f:
                    for i, line in enumerate(f):
                        t = _short(line)
                        if t:
                            lines.append(t)
                            if len(lines) >= max_lines:
                                break
            elif suf == ".docx":
                print(f"[docx] {p}")
                for t in f_scandocx(p, max_items=min(300, max_lines - len(lines))):
                    lines.append(t)
                    if len(lines) >= max_lines:
                        break
            elif suf == ".xlsx":
                print(f"[xlsx] {p}")
                for t in f_scanxlsx(p, max_items=min(400, max_lines - len(lines))):
                    lines.append(t)
                    if len(lines) >= max_lines:
                        break
            elif suf == ".xls":
                print(f"[xls] {p}")
                for t in f_scanxls(p, max_items=min(400, max_lines - len(lines))):
                    lines.append(t)
                    if len(lines) >= max_lines:
                        break
        except Exception:
            pass

        # Kill resource hog
        if len(lines) >= max_lines:
            break

    # Sample
    titles = [t for t in titles if t]
    random.shuffle(titles)
    random.shuffle(lines)

    return titles[:500], lines[:1000]

def f_picktopics(titles, lines, k=8):
    pool                    = list(set(titles + lines))
    random.shuffle(pool)
    return pool[:k]

def f_generatenamemarkov(mk: Markov, topic: str) -> str:
    seed                    = topic.split()[:3]
    seed_str                = " ".join(seed) if seed else None
    raw = mk.generate(max_words=6, seed=seed_str)
    if not raw:
        raw                 = topic
    base                    = re.sub(r"[^a-zA-Z0-9]+", "_", raw).strip("_").lower()
    if not base:
        base                = re.sub(r"[^a-z0-9]+", "_", topic.lower()).strip("_")[:20] or "decoy"
    ext                     = random.choice([".txt", ".csv"])
    name                    = f"{base}_{DATE_TAG}{ext}"
    name                    = re.sub(r'_{2,}', '_', name)[:FILENAME_MAX_CHARS]
    return name

def f_generatebodymarkov(mk: Markov, topic: str) -> str:
    seed                    = topic
    body                    = mk.wordcount(target_min=BODYMIN, target_max=BODYMAX,seed=seed)
    body                    = re.sub(r"\s+\.\s+", ". ", body).strip()
    if len(body.split()) < BODYMIN // 2:
        body                = f"Internal note: {topic}\n\nThis is a confidential document. " 
    footer                  = "\n\nNote: approvals required before accessing production."
    if not body.endswith("Note: approvals required before accessing production."):
        body                = body.rstrip() + footer
    return body

def f_generatefirstname():
    names = ["James", "Michael", "John", "Robert", "David", "William", "Richard", "Joseph", "Thomas", "Christopher", "Charles", "Daniel", "Matthew", "Anthony", "Mark", "Steven", "Donald", "Andrew", "Joshua", "Paul", "Kenneth", "Kevin", "Brian", "Timothy", "Ronald", "Jason", "George", "Edward", "Jeffrey", "Ryan", "Jacob", "Nicholas", "Gary", "Eric", "Jonathan", "Stephen", "Larry", "Justin", "Benjamin", "Scott", "Brandon", "Samuel", "Gregory", "Alexander", "Patrick", "Frank", "Jack", "Raymond", "Dennis", "Tyler", "Aaron", "Jerry", "Jose", "Nathan", "Adam", "Henry", "Zachary", "Douglas", "Peter", "Noah", "Kyle", "Ethan", "Christian", "Jeremy", "Keith", "Austin", "Sean", "Roger", "Terry", "Walter", "Dylan", "Gerald", "Carl", "Jordan", "Bryan", "Gabriel", "Jesse", "Harold", "Lawrence", "Logan", "Arthur", "Bruce", "Billy", "Elijah", "Joe", "Alan", "Juan", "Liam", "Willie", "Mason", "Albert", "Randy", "Wayne", "Vincent", "Lucas", "Caleb", "Luke", "Bobby", "Isaac", "Bradley", "Mary", "Patricia", "Jennifer", "Linda", "Elizabeth", "Barbara", "Susan", "Jessica", "Karen", "Sarah", "Lisa", "Nancy", "Sandra", "Ashley", "Emily", "Kimberly", "Betty", "Margaret", "Donna", "Michelle", "Carol", "Amanda", "Melissa", "Deborah", "Stephanie", "Rebecca", "Sharon", "Laura", "Cynthia", "Amy", "Kathleen", "Angela", "Dorothy", "Shirley", "Emma", "Brenda", "Nicole", "Pamela", "Samantha", "Anna", "Katherine", "Christine", "Debra", "Rachel", "Olivia", "Carolyn", "Maria", "Janet", "Heather", "Diane", "Catherine", "Julie", "Victoria", "Helen", "Joyce", "Lauren", "Kelly", "Christina", "Joan", "Judith", "Ruth", "Hannah", "Evelyn", "Andrea", "Virginia", "Megan", "Cheryl", "Jacqueline", "Madison", "Sophia", "Abigail", "Teresa", "Isabella", "Sara", "Janice", "Martha", "Gloria", "Kathryn", "Ann", "Charlotte", "Judy", "Amber", "Julia", "Grace", "Denise", "Danielle", "Natalie", "Alice", "Marilyn", "Diana", "Beverly", "Jean", "Brittany", "Theresa", "Frances", "Kayla", "Alexis", "Tiffany", "Lori", "Kathy"]
    return random.choice(names)

def f_generatepwd():
    alphabet                = string.ascii_letters + string.digits + "!@#$%^&*()-_=+"
    pwd                     = ''.join(secrets.choice(alphabet) for _ in range(12))    
    return pwd

def f_generatejuicyname(topic: str, prefer_style: str = "snake") -> str:
    kw                      = random.choice(JUICY_KEYWORDS)
    tpl                     = random.choice(JUICY_NAME)
    base                    = tpl.format(kw=kw, date=JUICY_DATE_TAG)
    if prefer_style == "snake":
        filename            = re.sub(r"[^a-zA-Z0-9]+", "_", base).strip("_")
    else:
        filename            = re.sub(r"[^a-zA-Z0-9]+", "-", base).strip("-")
    ext                     = random.choice(JUICY_EXT_PREFS)
    return f"{filename}{ext}"

def f_generatejuicybody(topic: str, ext: str = ".txt") -> str:
    if ext.lower() in {".csv", ".txt", ".md"}:
        rows                = []
        rows.append("account_id,username,password,role,notes")
        for i in range(random.randint(1, 134)):
            acc             = f"{1000 + i}"
            user            = f_generatefirstname()
            pwd             = f_generatepwd()
            role            = random.choice(["admin","svc_account","finance","ops","user"])
            note            = random.choice(["temporary", "migrated", "legacy", "no access"])
            rows.append(f"{acc},{user},{pwd},{role},{note}")
        body                = "\n".join(rows)
        body                += "\n\nSource topic: {topic}\nNote: approvals required before production."
        return body

    text                    = f'''Internal credentials index for: {topic}\n\naccount_id | username | password | role\n'''       
    for i in range(random.randint(1, 134)):
        text += f"{1000+i} | {f_generatefirstname()} | {f_generatepwd} | {random.choice(['admin','svc','ops'])}\n"
    text += "Note: approvals required before production."
    return text

def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Scanning: {DOCUMENTS_DIR}")
    titles, lines           = f_scandir(DOCUMENTS_DIR)
    print(f"Title:{len(titles)}\nLines:{len(lines)}")
    mk                      = Markov(order=2)
    mk.train(titles + lines)
    topics                  = f_picktopics(titles, lines, k=NUM_DECOYS)
    print(f"Topics:{topics}")
    manifest = {"generated_at": int(time.time()), "source_dir": str(DOCUMENTS_DIR), "items": []}
    for i, topic in enumerate(topics, 1):
        try:
            fname           = f_generatenamemarkov(mk, topic)
            print(f'File:{fname}')
            body            = f_generatebodymarkov(mk, topic)
            print(f'Body:\n{body[:200]}\n{"-"*50}\n')
        except Exception as e:
            print("Generation error:", e)
            continue
        is_juicy            = random.random() < JUICY_PROB
        if is_juicy:
            print('JUICY!!!!')
            juicy_fname     = f_generatejuicyname(topic, prefer_style=random.choice(["snake", "kebab"]))
            stem            = f_safestem(juicy_fname)
            ext             = Path(juicy_fname).suffix or random.choice(JUICY_EXT_PREFS)
            target_path     = OUTPUT_DIR / f"{stem}{ext}"
            juicy_body      = f_generatejuicybody(topic, ext=ext)
            file_body       = juicy_body
            file_ext        = ext
        else:
            stem            = f_safestem(fname)
            file_ext        = random.choices([".txt", ".docx"], weights=[1,2], k=1)[0]
            target_path     = OUTPUT_DIR / f"{stem}{file_ext}"
            file_body       = body
        if target_path.exists():
            target_path     = OUTPUT_DIR / f"{stem}_{i}{file_ext}"
        if file_ext == ".docx":
            f_savedocx(file_body, target_path, title=topic if not is_juicy else f"{topic}")
        else:
            target_path.write_text(file_body + "\n", encoding="utf-8")
        manifest["items"].append({"topic": topic, "file": target_path.name, "juicy": bool(is_juicy)})
    (OUTPUT_DIR / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"Done. Generated {len(manifest['items'])} files in {OUTPUT_DIR}")

if __name__ == "__main__":
    main()
